# backend/ingestion.py 
import io, os, re, csv, pandas as pd 
from typing import Tuple 
from PyPDF2 import PdfReader 
from docx import Document as DocxDocument 
from backend.storage import SessionLocal, Document 
from datetime import datetime 
import json 
 
# Simple PII regex patterns for demo. Extend for enterprise use. 
PII_PATTERNS = { 
    "EMAIL": r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", 
    "PHONE": r"\b(?:\+?\d{1,3}[-.\s]?)?(?:\d{3}[-.\s]?\d{3}[-.\s]?\d{4})\b", 
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b", 
    "PAN": r"\b[A-Z]{5}[0-9]{4}[A-Z]\b" 
} 
 
def mask_pii(text: str, mask_token: str = "[REDACTED]") -> Tuple[str, dict]: 
    """Mask PII and return masked text and a summary of masked items.""" 
    masked = text 
    summary = {} 
    for label, pattern in PII_PATTERNS.items(): 
        matches = re.findall(pattern, masked) 
        if matches: 
            summary[label] = len(matches) 
            masked = re.sub(pattern, mask_token, masked) 
    return masked, summary 
 
def parse_pdf(path: str) -> str: 
    reader = PdfReader(path) 
    pages = [] 
    for p in reader.pages: 
        try: 
            pages.append(p.extract_text() or "") 
        except Exception: 
            pages.append("") 
    return "\n".join(pages) 
 
def parse_docx(path: str) -> str: 
    doc = DocxDocument(path) 
    return "\n".join([p.text for p in doc.paragraphs]) 
 
def parse_txt(path: str) -> str: 
    with open(path, "r", encoding="utf-8", errors="ignore") as f: 
        return f.read() 
 
def parse_csv(path: str) -> str: 
    df = pd.read_csv(path, dtype=str, keep_default_na=False) 
    return df.to_csv(index=False) 
 
def parse_document(path: str) -> Tuple[str, dict]: 
    ext = os.path.splitext(path)[1].lower() 
    if ext == ".pdf": 
        text = parse_pdf(path) 
    elif ext in [".docx", ".doc"]: 
        text = parse_docx(path) 
    elif ext == ".csv": 
        text = parse_csv(path) 
    else: 
        text = parse_txt(path) 
    masked_text, pii_summary = mask_pii(text) 
    return masked_text, pii_summary 
 
def index_document(filename: str, text: str, metadata: dict = None): 
    db = SessionLocal() 
    try: 
        doc = Document(filename=filename, text=text, metadata=metadata or {}) 
        db.add(doc) 
        db.commit() 
        db.refresh(doc) 
        return doc.id 
    finally: 
        db.close() 
 
def get_text_by_id(doc_id: int): 
    db = SessionLocal() 
    try: 
        doc = db.query(Document).filter(Document.id == doc_id).first() 
        return doc.text if doc else None 
    finally: 
        db.close() 